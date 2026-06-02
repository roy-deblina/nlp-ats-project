# file_parser.py
# Advanced Resume & JD Parsing Engine
# Optimized for ATS semantic embedding pipelines

import re
import io
import fitz
import docx
from typing import List


# ==========================================================
# GLOBAL CLEANING CONFIG
# ==========================================================

MIN_LINE_LENGTH = 2

COMMON_HEADERS = {
    "resume",
    "curriculum vitae",
    "cv"
}

EXCESSIVE_SYMBOLS_PATTERN = r"[■●◆►▪▶♦•]"

MULTISPACE_PATTERN = r"\s+"

EMAIL_PATTERN = r"\b[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}\b"

PHONE_PATTERN = r"""
(?:\+?\d{1,3}[\s\-]?)?
(?:\(?\d{3}\)?[\s\-]?)?
\d{3}[\s\-]?\d{4}
"""

URL_PATTERN = r"(https?://\S+|www\.\S+)"

PAGE_NUMBER_PATTERN = r"^\s*page\s+\d+\s*$"

# ==========================================================
# TEXT NORMALIZATION
# ==========================================================

def normalize_text(text: str) -> str:

    if not text:
        return ""

    # ------------------------------------------------------
    # Lower noise symbols
    # ------------------------------------------------------

    text = re.sub(EXCESSIVE_SYMBOLS_PATTERN, " ", text)

    # ------------------------------------------------------
    # Normalize unicode quotes/dashes
    # ------------------------------------------------------

    replacements = {
        "\u2013": "-",
        "\u2014": "-",
        "\u2018": "'",
        "\u2019": "'",
        "\u201c": '"',
        "\u201d": '"',
        "\xa0": " ",
        "\t": " "
    }

    for old, new in replacements.items():
        text = text.replace(old, new)

    # ------------------------------------------------------
    # Remove repeated punctuation
    # ------------------------------------------------------

    text = re.sub(r"\.{2,}", ".", text)
    text = re.sub(r"\-{2,}", "-", text)
    text = re.sub(r"\_{2,}", "_", text)

    # ------------------------------------------------------
    # Remove excessive whitespace
    # ------------------------------------------------------

    text = re.sub(MULTISPACE_PATTERN, " ", text)

    return text.strip()


# ==========================================================
# STRUCTURAL CLEANING
# ==========================================================

def clean_structured_lines(lines: List[str]) -> List[str]:

    cleaned = []

    for line in lines:

        line = line.strip()

        if not line:
            continue

        # --------------------------------------------------
        # Remove very short garbage lines
        # --------------------------------------------------

        if len(line) < MIN_LINE_LENGTH:
            continue

        # --------------------------------------------------
        # Remove page numbers
        # --------------------------------------------------

        if re.match(PAGE_NUMBER_PATTERN, line.lower()):
            continue

        # --------------------------------------------------
        # Remove common useless headers
        # --------------------------------------------------

        if line.lower() in COMMON_HEADERS:
            continue

        # --------------------------------------------------
        # Remove repeated separators
        # --------------------------------------------------

        if re.fullmatch(r"[-_=]{3,}", line):
            continue

        cleaned.append(line)

    return cleaned


# ==========================================================
# SECTION PRESERVATION
# ==========================================================

def preserve_sections(text: str) -> str:

    section_patterns = [
        "education",
        "experience",
        "skills",
        "projects",
        "certifications",
        "summary",
        "technical skills",
        "work experience",
        "professional experience",
        "achievements",
        "internships"
    ]

    for pattern in section_patterns:

        regex = rf"\b({pattern})\b"

        text = re.sub(
            regex,
            r"\n\n\1\n",
            text,
            flags=re.IGNORECASE
        )

    return text


# ==========================================================
# ENTITY PRESERVATION
# ==========================================================

def preserve_important_entities(text: str) -> str:

    # ------------------------------------------------------
    # Normalize emails
    # ------------------------------------------------------

    text = re.sub(
        EMAIL_PATTERN,
        lambda m: m.group(0).lower(),
        text
    )

    # ------------------------------------------------------
    # Normalize URLs
    # ------------------------------------------------------

    text = re.sub(
        URL_PATTERN,
        lambda m: m.group(0).lower(),
        text
    )

    # ------------------------------------------------------
    # Normalize phone spacing
    # ------------------------------------------------------

    text = re.sub(
        PHONE_PATTERN,
        lambda m: re.sub(r"\s+", "", m.group(0)),
        text,
        flags=re.VERBOSE
    )

    return text


# ==========================================================
# FINAL ATS CLEANING PIPELINE
# ==========================================================

def clean_document_text(text: str) -> str:

    if not text:
        return ""

    # ------------------------------------------------------
    # Normalize raw text
    # ------------------------------------------------------

    text = normalize_text(text)

    # ------------------------------------------------------
    # Split into structured lines
    # ------------------------------------------------------

    lines = text.splitlines()

    # ------------------------------------------------------
    # Remove junk lines
    # ------------------------------------------------------

    lines = clean_structured_lines(lines)

    # ------------------------------------------------------
    # Rejoin clean text
    # ------------------------------------------------------

    text = "\n".join(lines)

    # ------------------------------------------------------
    # Preserve semantic sections
    # ------------------------------------------------------

    text = preserve_sections(text)

    # ------------------------------------------------------
    # Preserve ATS-important entities
    # ------------------------------------------------------

    text = preserve_important_entities(text)

    # ------------------------------------------------------
    # Remove excessive blank lines
    # ------------------------------------------------------

    text = re.sub(r"\n{3,}", "\n\n", text)

    return text.strip()


# ==========================================================
# PDF PARSER
# ==========================================================

def extract_text_from_pdf(file_bytes: bytes) -> str:

    extracted_text = []

    try:

        pdf_document = fitz.open(
            stream=file_bytes,
            filetype="pdf"
        )

        for page in pdf_document:

            # --------------------------------------------------
            # Better structured extraction
            # --------------------------------------------------

            text = page.get_text(
                "text",
                sort=True
            )

            if text:
                extracted_text.append(text)

        pdf_document.close()

        final_text = "\n".join(extracted_text)

        return clean_document_text(final_text)

    except Exception as e:

        return f"""
        PDF Parsing Error:
        {str(e)}
        """


# ==========================================================
# DOCX PARSER
# ==========================================================

def extract_text_from_docx(file_bytes: bytes) -> str:

    extracted_text = []

    try:

        document = docx.Document(
            io.BytesIO(file_bytes)
        )

        # --------------------------------------------------
        # Extract paragraphs
        # --------------------------------------------------

        for para in document.paragraphs:

            text = para.text.strip()

            if text:
                extracted_text.append(text)

        # --------------------------------------------------
        # Extract tables
        # --------------------------------------------------

        for table in document.tables:

            for row in table.rows:

                row_data = []

                for cell in row.cells:

                    cell_text = cell.text.strip()

                    if cell_text:
                        row_data.append(cell_text)

                if row_data:
                    extracted_text.append(
                        " | ".join(row_data)
                    )

        final_text = "\n".join(extracted_text)

        return clean_document_text(final_text)

    except Exception as e:

        return f"""
        DOCX Parsing Error:
        {str(e)}
        """


# ==========================================================
# TXT PARSER
# ==========================================================

def extract_text_from_txt(file_bytes: bytes) -> str:

    try:

        text = file_bytes.decode(
            "utf-8",
            errors="ignore"
        )

        return clean_document_text(text)

    except Exception as e:

        return f"""
        TXT Parsing Error:
        {str(e)}
        """


# ==========================================================
# MAIN FILE ROUTER
# ==========================================================

def parse_uploaded_file(uploaded_file) -> str:

    if uploaded_file is None:
        return ""

    try:

        file_name = uploaded_file.name.lower()

        file_bytes = uploaded_file.read()

        # --------------------------------------------------
        # PDF
        # --------------------------------------------------

        if file_name.endswith(".pdf"):

            return extract_text_from_pdf(file_bytes)

        # --------------------------------------------------
        # DOCX
        # --------------------------------------------------

        elif file_name.endswith(".docx"):

            return extract_text_from_docx(file_bytes)

        # --------------------------------------------------
        # TXT
        # --------------------------------------------------

        elif file_name.endswith(".txt"):

            return extract_text_from_txt(file_bytes)

        # --------------------------------------------------
        # UNSUPPORTED
        # --------------------------------------------------

        else:

            return f"""
            Unsupported file format:
            {file_name}
            """

    except Exception as e:

        return f"""
        File Parsing Error:
        {str(e)}
        """